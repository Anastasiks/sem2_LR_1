import docx
from time import *
from docx import Document
import Levenshtein
from fuzzywuzzy import fuzz
def get_text_from_docx(path):
    doc = Document(path)
    return "\n".join([para.text for para in doc.paragraphs])


file1 = "docx1.docx"
file2 = "docx2.docx"

text1 = get_text_from_docx(file1)
text2 = get_text_from_docx(file2)

time1 = time()
distance = Levenshtein.distance(text1, text2)
time2 = time()

similarity = Levenshtein.ratio(text1, text2)*100

print(f"Расстояние Левенштейна: {distance}")
print(f"сходство в процентах {similarity}")
print(f"прошло времени {time2 - time1}")




def get_text(path):
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

text1 = get_text("docx1.docx")
text2 = get_text("docx2.docx")
time3 = time()
ratio = fuzz.ratio(text1, text2)
time4 = time()
print(f"fuzzy: {ratio}")
print(f"за  {time4 - time3}")

def levenstein(text1, text2):
    n, m = len(text1), len(text2)
    if n > m:
        text1, text2 = text2, text1
        n, m = m, n

    current_row = range(n + 1)
    for i in range(1, m + 1):
        previous_row, current_row = current_row, [i] + [0] * n
        for j in range(1, n + 1):
            add, delete, change = previous_row[j] + 1, current_row[j - 1] + 1, previous_row[j - 1]
            if text1[j - 1] != text2[i - 1]:
                change += 1
            current_row[j] = min(add, delete, change)

    return current_row[n]
time5 = time()
print(f"через дп {levenstein(text1, text2)}")
time6 = time()
print(f"за {time6 - time5}")